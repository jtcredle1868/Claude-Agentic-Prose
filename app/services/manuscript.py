"""Service for generating industry-standard manuscript submission packets."""

import os
import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors

from app.services.ai_client import AIClient


class ManuscriptService:
    """Generate complete manuscript submission packets with industry-standard formatting."""

    STANDARD_FONT = "Times New Roman"
    STANDARD_SIZE = 12
    STANDARD_MARGINS = 1.0  # inches
    LINES_PER_PAGE_ESTIMATE = 25
    WORDS_PER_PAGE_ESTIMATE = 250

    def __init__(self, ai_client=None):
        self.ai = ai_client or AIClient()

    # ─── DOCX Manuscript Generation ──────────────────────────────────────

    def generate_full_manuscript_docx(self, project, output_path):
        """Generate a complete, properly formatted manuscript DOCX."""
        doc = Document()
        self._setup_manuscript_styles(doc)
        self._add_title_page(doc, project)
        self._add_manuscript_body(doc, project)
        doc.save(output_path)
        return output_path

    def _setup_manuscript_styles(self, doc):
        """Set up standard manuscript formatting."""
        style = doc.styles["Normal"]
        font = style.font
        font.name = self.STANDARD_FONT
        font.size = Pt(self.STANDARD_SIZE)
        pf = style.paragraph_format
        pf.line_spacing = 2.0
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        pf.first_line_indent = Inches(0.5)

        for section in doc.sections:
            section.top_margin = Inches(self.STANDARD_MARGINS)
            section.bottom_margin = Inches(self.STANDARD_MARGINS)
            section.left_margin = Inches(self.STANDARD_MARGINS)
            section.right_margin = Inches(self.STANDARD_MARGINS)

    def _add_title_page(self, doc, project):
        """Add standard manuscript title page."""
        # Top left: author contact info
        word_count = project.word_count
        rounded_count = round(word_count / 1000) * 1000 if word_count > 1000 else word_count

        contact = doc.add_paragraph()
        contact.paragraph_format.first_line_indent = Inches(0)
        contact.paragraph_format.line_spacing = 1.0
        contact_lines = [project.author_name or "Author Name"]
        if project.author_address:
            contact_lines.append(project.author_address)
        if project.author_email:
            contact_lines.append(project.author_email)
        if project.author_phone:
            contact_lines.append(project.author_phone)
        if project.agent_name:
            contact_lines.append(f"Agent: {project.agent_name}")
        contact.text = "\n".join(contact_lines)

        # Word count top right (approximated via right-aligned paragraph)
        wc_para = doc.add_paragraph()
        wc_para.paragraph_format.first_line_indent = Inches(0)
        wc_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        wc_para.text = f"Approx. {rounded_count:,} words"

        # Vertical space to center title
        for _ in range(10):
            spacer = doc.add_paragraph()
            spacer.paragraph_format.first_line_indent = Inches(0)

        # Title centered
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.first_line_indent = Inches(0)
        run = title_para.add_run(project.title.upper())
        run.font.size = Pt(self.STANDARD_SIZE)
        run.font.name = self.STANDARD_FONT

        # Subtitle if present
        if project.subtitle:
            sub_para = doc.add_paragraph()
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_para.paragraph_format.first_line_indent = Inches(0)
            sub_para.text = project.subtitle

        # "by" line
        by_para = doc.add_paragraph()
        by_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        by_para.paragraph_format.first_line_indent = Inches(0)
        by_para.text = f"by\n{project.author_name or 'Author Name'}"

        # Page break after title
        doc.add_page_break()

    def _add_manuscript_body(self, doc, project):
        """Add all chapters with proper manuscript formatting."""
        chapters = sorted(project.chapters, key=lambda c: c.order)
        author_last = (project.author_name or "Author").split()[-1]
        short_title = project.title[:30].upper() if project.title else "UNTITLED"

        for i, chapter in enumerate(chapters):
            # Chapter heading: start 1/3 down the page
            for _ in range(4):
                spacer = doc.add_paragraph()
                spacer.paragraph_format.first_line_indent = Inches(0)

            # Chapter title centered
            ch_title = doc.add_paragraph()
            ch_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ch_title.paragraph_format.first_line_indent = Inches(0)
            ch_title.paragraph_format.space_after = Pt(24)
            run = ch_title.add_run(f"Chapter {chapter.order}: {chapter.title}")
            run.font.size = Pt(self.STANDARD_SIZE)
            run.font.name = self.STANDARD_FONT

            # Chapter content
            content = chapter.content or ""
            paragraphs = content.split("\n")
            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text:
                    continue
                if para_text == "###" or para_text == "* * *":
                    # Scene break
                    scene_break = doc.add_paragraph()
                    scene_break.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    scene_break.paragraph_format.first_line_indent = Inches(0)
                    scene_break.text = "#"
                    continue
                p = doc.add_paragraph(para_text)

            # Page break between chapters (not after last)
            if i < len(chapters) - 1:
                doc.add_page_break()

        # End marker
        end_para = doc.add_paragraph()
        end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        end_para.paragraph_format.first_line_indent = Inches(0)
        end_para.paragraph_format.space_before = Pt(48)
        end_para.text = "# # #"

    # ─── Query Letter Generation ─────────────────────────────────────────

    def generate_query_letter(self, project):
        """Generate an industry-standard query letter."""
        system = (
            "You are a literary agent's assistant who has read thousands of successful "
            "query letters. You write compelling, professionally formatted query letters "
            "that follow industry standards precisely."
        )
        prompt = f"""Write a professional query letter for the following project:

TITLE: {project.title}
GENRE: {project.genre}
TYPE: {project.project_type}
WORD COUNT: {project.word_count:,}
SYNOPSIS: {project.synopsis}
AUTHOR NAME: {project.author_name or 'Author Name'}
AUTHOR BIO: {project.author_bio or 'Provide relevant biographical details'}

Format the query letter following standard industry conventions:
1. Opening hook (1-2 sentences that grab attention)
2. Book description (2-3 paragraphs covering protagonist, conflict, stakes, and unique elements)
3. Comparable titles (2-3 comp titles with brief explanation)
4. Author bio paragraph (relevant credentials and platform)
5. Professional closing with word count, genre, and title

The letter should be under 400 words total. Make it compelling and specific—
avoid generic language. Every sentence should earn its place."""

        return self.ai.generate(system, prompt, max_tokens=2048)

    # ─── Synopsis Generation ─────────────────────────────────────────────

    def generate_synopsis(self, project, length="standard"):
        """Generate a manuscript synopsis in various lengths."""
        lengths = {
            "short": "1-2 pages (approximately 500 words)",
            "standard": "3-5 pages (approximately 1,500 words)",
            "detailed": "8-10 pages (approximately 3,000 words)",
        }
        target = lengths.get(length, lengths["standard"])

        chapters_summary = ""
        for ch in sorted(project.chapters, key=lambda c: c.order):
            chapters_summary += f"\nChapter {ch.order}: {ch.title}\n{ch.summary}\n"

        system = (
            "You are an expert at writing manuscript synopses that agents and editors "
            "actually want to read. You distill complex narratives into compelling, "
            "clear summaries that reveal the full arc including the ending."
        )
        prompt = f"""Write a {target} synopsis for:

TITLE: {project.title}
GENRE: {project.genre}
TYPE: {project.project_type}

CHAPTER SUMMARIES:
{chapters_summary}

THEMES: {project.themes}

Requirements:
- Written in present tense, third person
- Include ALL major plot points including the ending
- Introduce main characters with their full names in CAPS on first appearance
- Show the emotional arc alongside the plot arc
- Maintain the tone of the manuscript
- Use paragraph breaks for readability
- Do NOT include chapter divisions—this should read as a continuous narrative summary"""

        return self.ai.generate(system, prompt, max_tokens=8192)

    # ─── Full Submission Packet ──────────────────────────────────────────

    def generate_submission_packet(self, project, output_dir):
        """Generate a complete submission packet with all standard documents."""
        os.makedirs(output_dir, exist_ok=True)
        generated_files = []

        # 1. Full manuscript DOCX
        manuscript_path = os.path.join(output_dir, f"{self._safe_filename(project.title)}_manuscript.docx")
        self.generate_full_manuscript_docx(project, manuscript_path)
        generated_files.append(("Full Manuscript", manuscript_path))

        # 2. Query letter
        query_text = self.generate_query_letter(project)
        query_path = os.path.join(output_dir, f"{self._safe_filename(project.title)}_query_letter.docx")
        self._text_to_docx(query_text, query_path, "Query Letter")
        generated_files.append(("Query Letter", query_path))

        # 3. Synopsis (short and standard)
        for syn_length in ["short", "standard"]:
            synopsis_text = self.generate_synopsis(project, length=syn_length)
            syn_path = os.path.join(output_dir, f"{self._safe_filename(project.title)}_synopsis_{syn_length}.docx")
            self._text_to_docx(synopsis_text, syn_path, f"Synopsis ({syn_length.title()})")
            generated_files.append((f"Synopsis ({syn_length.title()})", syn_path))

        # 4. Author bio
        bio_text = self._generate_author_bio(project)
        bio_path = os.path.join(output_dir, f"{self._safe_filename(project.title)}_author_bio.docx")
        self._text_to_docx(bio_text, bio_path, "Author Biography")
        generated_files.append(("Author Biography", bio_path))

        # 5. Chapter-by-chapter outline
        outline_text = self._generate_chapter_outline(project)
        outline_path = os.path.join(output_dir, f"{self._safe_filename(project.title)}_chapter_outline.docx")
        self._text_to_docx(outline_text, outline_path, "Chapter-by-Chapter Outline")
        generated_files.append(("Chapter Outline", outline_path))

        # 6. First three chapters (sample pages)
        sample_path = os.path.join(output_dir, f"{self._safe_filename(project.title)}_sample_chapters.docx")
        self._generate_sample_chapters(project, sample_path)
        generated_files.append(("Sample Chapters", sample_path))

        return generated_files

    def _generate_author_bio(self, project):
        """Generate author bio document."""
        if project.author_bio:
            return project.author_bio

        system = (
            "You are helping an author draft their professional biography for "
            "manuscript submissions."
        )
        prompt = f"""Draft a professional author biography for a {project.genre} {project.project_type} writer.
Author name: {project.author_name or 'Author Name'}
Available info: {project.author_bio or 'No details provided yet.'}

Write a 150-200 word professional bio in third person that:
- Establishes credibility for the genre
- Mentions relevant education, experience, or expertise
- Includes a personal detail or two for warmth
- Notes any previous publications
- Ends with where the author lives

Mark any placeholder information with [PLACEHOLDER] for the author to fill in."""

        return self.ai.generate(system, prompt, max_tokens=1024)

    def _generate_chapter_outline(self, project):
        """Generate a chapter-by-chapter outline document."""
        lines = [f"CHAPTER-BY-CHAPTER OUTLINE", f"{project.title}", f"by {project.author_name or 'Author'}", ""]
        for ch in sorted(project.chapters, key=lambda c: c.order):
            lines.append(f"Chapter {ch.order}: {ch.title}")
            lines.append(ch.summary or "[Summary to be added]")
            lines.append(f"Word count: {ch.word_count:,}")
            lines.append("")
        lines.append(f"Total word count: {project.word_count:,}")
        return "\n".join(lines)

    def _generate_sample_chapters(self, project, output_path):
        """Generate a sample chapters document (first three chapters)."""
        doc = Document()
        self._setup_manuscript_styles(doc)
        self._add_title_page(doc, project)

        chapters = sorted(project.chapters, key=lambda c: c.order)[:3]
        for i, chapter in enumerate(chapters):
            for _ in range(4):
                spacer = doc.add_paragraph()
                spacer.paragraph_format.first_line_indent = Inches(0)

            ch_title = doc.add_paragraph()
            ch_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ch_title.paragraph_format.first_line_indent = Inches(0)
            ch_title.paragraph_format.space_after = Pt(24)
            run = ch_title.add_run(f"Chapter {chapter.order}: {chapter.title}")
            run.font.size = Pt(self.STANDARD_SIZE)
            run.font.name = self.STANDARD_FONT

            content = chapter.content or "[Chapter content to be added]"
            for para_text in content.split("\n"):
                para_text = para_text.strip()
                if not para_text:
                    continue
                if para_text in ("###", "* * *"):
                    scene_break = doc.add_paragraph()
                    scene_break.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    scene_break.paragraph_format.first_line_indent = Inches(0)
                    scene_break.text = "#"
                    continue
                doc.add_paragraph(para_text)

            if i < len(chapters) - 1:
                doc.add_page_break()

        doc.save(output_path)

    def _text_to_docx(self, text, output_path, title=""):
        """Convert plain text to a formatted DOCX document."""
        doc = Document()
        self._setup_manuscript_styles(doc)

        if title:
            t = doc.add_paragraph()
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            t.paragraph_format.first_line_indent = Inches(0)
            t.paragraph_format.space_after = Pt(24)
            run = t.add_run(title)
            run.font.size = Pt(self.STANDARD_SIZE)
            run.font.name = self.STANDARD_FONT
            run.bold = True

        for para_text in text.split("\n"):
            if not para_text.strip():
                doc.add_paragraph("")
                continue
            p = doc.add_paragraph(para_text.strip())
            p.paragraph_format.first_line_indent = Inches(0)

        doc.save(output_path)

    @staticmethod
    def _safe_filename(name):
        """Create a filesystem-safe filename."""
        safe = "".join(c for c in name if c.isalnum() or c in (" ", "-", "_")).strip()
        return safe.replace(" ", "_")[:50] or "untitled"
